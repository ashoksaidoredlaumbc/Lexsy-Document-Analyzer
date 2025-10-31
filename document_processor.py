import re
from docx import Document
from typing import List, Dict
import mammoth

class DocumentProcessor:
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        
    def extract_text(self) -> str:
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text)
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return "\n".join(full_text)
    
    def detect_placeholders(self) -> List[str]:
        text = self.extract_text()
        
        patterns = [
            r'\{([^}]+)\}',           # {placeholder}
            r'\[([^\]]+)\]',          # [PLACEHOLDER]
            r'\{\{([^}]+)\}\}',       # {{placeholder}}
            r'__+([A-Z_\s]+)__+',     # ___PLACEHOLDER___
            r'\$\{([^}]+)\}',         # ${placeholder}
        ]
        
        placeholders = set()
        for pattern in patterns:
            matches = re.findall(pattern, text)
            placeholders.update([m.strip() for m in matches])
        
        filtered = [p for p in placeholders if len(p) > 1 and not p.isdigit()]
        print(f"Detected placeholders: {filtered}") 
        
        return sorted(list(set(filtered)))


    def fill_template(self, placeholder_values: Dict[str, str], output_path: str):
        """Replace placeholders with actual values - reload fresh document each time"""
        
        doc = Document(self.file_path)
        
        # Process paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            modified_text = original_text
            
            for key, value in placeholder_values.items():
                formats = [
                    f"{{{key}}}",           # {key}
                    f"[{key}]",             # [key]
                    f"{{{{{key}}}}}",       # {{key}}
                    f"${{{key}}}",          # ${key}
                    f"$[{key}]",            # $[key]
                    key                     # raw key
                ]
                
                for fmt in formats:
                    if fmt in modified_text:
                        modified_text = modified_text.replace(fmt, str(value))
            
            if modified_text != original_text:
                para.text = modified_text
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    modified_text = original_text
                    
                    for key, value in placeholder_values.items():
                        formats = [
                            f"{{{key}}}",
                            f"[{key}]",
                            f"{{{{{key}}}}}",
                            f"${{{key}}}",
                            f"$[{key}]",
                            key
                        ]
                        
                        for fmt in formats:
                            if fmt in modified_text:
                                modified_text = modified_text.replace(fmt, str(value))
                    
                    if modified_text != original_text:
                        cell.text = modified_text
   
        doc.save(output_path)
        print(f"Document saved to: {output_path}")  
        print(f"Values used: {placeholder_values}")


    def generate_html_preview(self, file_path: str) -> str:
        """Convert DOCX to HTML for preview"""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
        
                styled_html = f"""
                <div class="document-preview">
                    <style>
                        .document-preview {{
                            font-family: 'Times New Roman', serif;
                            line-height: 1.6;
                            padding: 40px;
                            background: white;
                            max-width: 800px;
                            margin: 0 auto;
                            box-shadow: 0 0 10px rgba(0,0,0,0.1);
                        }}
                        .document-preview p {{
                            margin-bottom: 12px;
                        }}
                        .document-preview table {{
                            width: 100%;
                            border-collapse: collapse;
                            margin: 20px 0;
                        }}
                        .document-preview table td, .document-preview table th {{
                            border: 1px solid #ddd;
                            padding: 8px;
                        }}
                    </style>
                    {html_content}
                </div>
                """
                return styled_html
        except Exception as e:
            return f"<p>Error generating preview: {str(e)}</p>"

    def get_placeholder_context(self, placeholder: str, context_chars: int = 300) -> str:
        text = self.extract_text()

        formats = [
            f"{{{placeholder}}}",
            f"[{placeholder}]",
            f"{{{{{placeholder}}}}}",
            f"${{{placeholder}}}",
            placeholder 
        ]

        for fmt in formats:
            if fmt in text:
                pos = text.find(fmt)
                
                start = max(0, pos - context_chars)
                end = min(len(text), pos + len(fmt) + context_chars)
                
                context = text[start:end]
                
                context = context.replace('\n', ' ').strip()
                
                return context

        return ""

